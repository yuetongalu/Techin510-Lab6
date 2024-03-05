from tempfile import NamedTemporaryFile
import os

import streamlit as st
from spellchecker import SpellChecker  # For spelling corrections
from textblob import TextBlob  # For grammar and sentiment analysis
from dotenv import load_dotenv
from pdfminer.high_level import extract_text  # For extracting text from PDF

from time import sleep

# Assuming PDF and possibly DOCX processing capabilities are desired
# Import necessary libraries for PDF and DOCX processing if available
try:
    from pdfminer.high_level import extract_text as extract_text_pdf
except ImportError:
    def extract_text_pdf(file_path):
        raise NotImplementedError("PDF processing not available.")

try:
    from docx import Document
except ImportError:
    def Document(file_path):
        raise NotImplementedError("DOCX processing not available.")

def extract_text_docx(file_path):
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# Streamlit App Configuration
st.set_page_config(
    page_title="Resume Fixer",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="auto",
)

# File Upload
uploaded_file = st.file_uploader("Upload your resume", type=['pdf', 'docx'])

if uploaded_file:
    bytes_data = uploaded_file.read()
    
    with NamedTemporaryFile(delete=False) as tmp:
        tmp.write(bytes_data)
        tmp_path = tmp.name
    
    try:
        if uploaded_file.type == "application/pdf":
            text = extract_text_pdf(tmp_path)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_docx(tmp_path)
        else:
            st.error("Unsupported file type.")
            text = ""
    finally:
        # Attempt to remove the temporary file with retries
        for attempt in range(3):
            try:
                os.remove(tmp_path)
                break
            except PermissionError:
                sleep(1)  # Wait for 1 second before retrying
        else:
            st.warning("Could not delete temporary file, it might still be locked.")

    if text:
        with st.expander("Extracted Text"):
            st.write(text)

    if text:
        with st.expander("Original Resume Text"):
            st.write(text)
        
        # Example analysis (extend with your own logic)
        spell = SpellChecker()
        blob = TextBlob(text)
        
        # Find misspelled words
        words = text.split()
        misspelled = spell.unknown(words)
        
        corrected_text = text
        for word in misspelled:
            # Suggest correction for each misspelled word
            corrected_text = corrected_text.replace(word, spell.correction(word))
        
        grammar_feedback = "Grammar looks good!" if blob.correct() == text else "There might be some grammatical errors."
        
        with st.expander("Corrected Resume Text"):
            st.write(corrected_text)
        
        with st.expander("Feedback"):
            st.write(f"Spelling Corrections: {', '.join(misspelled)}")
            st.write(f"Grammar Feedback: {grammar_feedback}")
